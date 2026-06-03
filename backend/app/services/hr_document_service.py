"""
HR letter PDF generation.
"""

from __future__ import annotations

import io
import re
from datetime import date
from pathlib import Path

from PIL import Image as PILImage
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.colors import HexColor, black
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas
from reportlab.platypus import Paragraph

from app.schemas.hr_document import InternshipCompletionLetterRequest

ASSETS_DIR = Path(__file__).resolve().parents[1] / "certificate_assets"
FONTS_DIR = ASSETS_DIR / "fonts"
IMAGES_DIR = ASSETS_DIR / "images"
LOGO_ICON_PATH = IMAGES_DIR / "reknew_logo_icon.png"
WORDMARK_PATH = IMAGES_DIR / "reknew_wordmark_official.png"
SIGNATURE_PATH = IMAGES_DIR / "signature.png"
SIGNATURE_CLEAN_PATH = IMAGES_DIR / "signature_clean.png"
SIGNATURE_OFFICIAL_PATH = IMAGES_DIR / "murali_sajja_signature_official.png"
_fonts_registered = False


def register_document_fonts() -> None:
    global _fonts_registered
    if _fonts_registered:
        return

    font_specs = {
        "Poppins": "Poppins-Regular.ttf",
        "Poppins-Bold": "Poppins-Bold.ttf",
        "Bricolage-Bold": "BricolageGrotesque-Bold.ttf",
    }
    for name, filename in font_specs.items():
        path = FONTS_DIR / filename
        if path.exists():
            try:
                pdfmetrics.registerFont(TTFont(name, str(path)))
            except Exception:
                pass
    _fonts_registered = True


def font(preferred: str, fallback: str = "Helvetica") -> str:
    try:
        pdfmetrics.getFont(preferred)
        return preferred
    except Exception:
        return fallback


def format_document_date(value: date) -> str:
    return f"{value.strftime('%B')} {value.day}, {value.year}"


def clean_filename(value: str) -> str:
    return re.sub(r"[^A-Za-z0-9_-]+", "_", value.strip()).strip("_")


def build_internship_completion_filename(intern_name: str, extension: str = "pdf") -> str:
    return f"Internship_Completion_Letter_{clean_filename(intern_name) or 'Intern'}.{extension}"


def pil_to_reader(image: PILImage.Image) -> ImageReader:
    buffer = io.BytesIO()
    image.save(buffer, format="PNG")
    buffer.seek(0)
    return ImageReader(buffer)


def draw_image(c: canvas.Canvas, path: Path, x: float, y: float, w: float, h: float, opacity: float = 1) -> None:
    if not path.exists():
        return

    image = PILImage.open(path).convert("RGBA")
    if opacity < 1:
        alpha = image.getchannel("A")
        alpha = alpha.point(lambda pixel: int(pixel * opacity))
        image.putalpha(alpha)

    c.drawImage(pil_to_reader(image), x, y, width=w, height=h, mask="auto")


def draw_wordmark(c: canvas.Canvas, right_x: float, y: float) -> None:
    if WORDMARK_PATH.exists():
        width = 100
        height = 24
        draw_image(c, WORDMARK_PATH, right_x - width, y - 4, width, height)
        return

    word_font = font("Bricolage-Bold", "Helvetica-Bold")
    size = 22
    icon_size = 12
    gap = 1.2
    re_w = pdfmetrics.stringWidth("re", word_font, size)
    knew_w = pdfmetrics.stringWidth("knew", word_font, size)
    total_w = re_w + icon_size + knew_w + gap * 2
    x = right_x - total_w

    c.setFillColor(HexColor("#202938"))
    c.setFont(word_font, size)
    c.drawString(x, y, "re")
    draw_image(c, LOGO_ICON_PATH, x + re_w + gap, y + 4.5, icon_size, icon_size)
    c.drawString(x + re_w + gap + icon_size + gap, y, "knew")


def draw_signature(c: canvas.Canvas, x: float, y: float) -> None:
    signature_path = (
        SIGNATURE_OFFICIAL_PATH
        if SIGNATURE_OFFICIAL_PATH.exists()
        else SIGNATURE_CLEAN_PATH
        if SIGNATURE_CLEAN_PATH.exists()
        else SIGNATURE_PATH
    )
    draw_image(c, signature_path, x, y, 126, 36)


def draw_wrapped_paragraph(
    c: canvas.Canvas,
    text: str,
    x: float,
    y: float,
    width: float,
    style: ParagraphStyle,
) -> float:
    paragraph = Paragraph(text, style)
    _, height = paragraph.wrap(width, 700)
    paragraph.drawOn(c, x, y - height)
    return y - height


def intern_first_name(intern_name: str) -> str:
    return intern_name.strip().split()[0] if intern_name.strip() else "the intern"


def normalize_responsibility_summary(intern_name: str, summary: str) -> str:
    value = summary.strip()
    if not value:
        return value
    first = intern_first_name(intern_name)
    if value[:1].islower() or value.lower().startswith(("was ", "worked ", "participated ", "contributed ")):
        return f"{first} {value}"
    return value


def internship_completion_body_parts(request: InternshipCompletionLetterRequest) -> list[str]:
    first = intern_first_name(request.intern_name)
    responsibility = normalize_responsibility_summary(request.intern_name, request.responsibility_summary)
    start = format_document_date(request.start_date)
    end = format_document_date(request.end_date)
    return [
        "To whom it may concern,",
        (
            f"We are pleased to confirm that {request.intern_name.strip()} successfully completed their internship "
            f"with ReKnew from {start} to {end}."
        ),
        f"During the internship, {responsibility}",
        (
            "At ReKnew, we believe in providing practical exposure and meaningful learning opportunities, and "
            f"{first} embraced those opportunities with dedication and enthusiasm."
        ),
        (
            "We sincerely appreciate the effort and positive contribution made during the internship and are "
            "confident that the experience gained here will support future academic and professional success."
        ),
        f"We wish {first} all the very best in future endeavors.",
        "Sincerely,",
    ]


def generate_internship_completion_pdf(request: InternshipCompletionLetterRequest) -> bytes:
    register_document_fonts()

    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    page_w, page_h = letter

    left = 1.0 * inch
    right = page_w - 1.0 * inch
    width = right - left
    top = page_h - 0.55 * inch
    body_font = font("Poppins", "Helvetica")
    body_bold = font("Poppins-Bold", "Helvetica-Bold")
    title_font = font("Poppins-Bold", "Helvetica-Bold")

    draw_wordmark(c, right, top - 8)

    c.setFillColor(black)
    c.setFont(title_font, 15)
    title = "Internship Completion Letter"
    c.drawCentredString(page_w / 2, page_h - 1.15 * inch, title)

    style = ParagraphStyle(
        "LetterBody",
        fontName=body_font,
        fontSize=10.8,
        leading=16.8,
        textColor=HexColor("#171717"),
        spaceAfter=0,
    )
    small_style = ParagraphStyle(
        "Footer",
        fontName=body_font,
        fontSize=7.5,
        leading=10,
        textColor=HexColor("#8A8A8A"),
    )

    y = page_h - 1.55 * inch
    c.setFont(body_font, 10.4)
    c.drawString(left, y, format_document_date(request.issued_date))
    y -= 0.37 * inch

    company = (
        f'<font name="{body_bold}">ReKnew Business Solutions Inc.</font><br/>'
        "7800 N. Dallas Pkwy, Ste. 320<br/>"
        "Plano, TX 75024"
    )
    y = draw_wrapped_paragraph(c, company, left, y, width, style) - 0.30 * inch

    body_parts = internship_completion_body_parts(request)

    for paragraph in body_parts:
        y = draw_wrapped_paragraph(c, paragraph, left, y, width, style) - 0.20 * inch

    y -= 0.04 * inch
    draw_signature(c, left - 4, y - 0.32 * inch)
    y -= 0.52 * inch

    c.setFont(body_bold, 10)
    c.setFillColor(black)
    c.drawString(left, y, "Murali Sajja | CEO")
    y -= 14
    c.drawString(left, y, "ReKnew | reknew.ai")

    footer_y = 0.77 * inch
    draw_image(c, LOGO_ICON_PATH, page_w - 4.20 * inch, -0.18 * inch, 3.85 * inch, 2.85 * inch, opacity=0.045)
    c.setStrokeColor(black)
    c.setLineWidth(0.6)
    c.line(left, footer_y + 0.14 * inch, right, footer_y + 0.14 * inch)
    draw_wrapped_paragraph(c, "© ReKnew. All Rights Reserved.", left, footer_y - 0.02 * inch, width / 2, small_style)
    c.setFont(body_font, 7.5)
    c.setFillColor(HexColor("#8A8A8A"))
    c.drawRightString(right, footer_y - 0.02 * inch, "www.reknew.ai")

    c.save()
    return buffer.getvalue()


def set_run_font(run, bold: bool = False, size: float = 11, color: str = "171717") -> None:
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_docx_paragraph(document: Document, text: str, after: float = 12):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def generate_internship_completion_docx(request: InternshipCompletionLetterRequest) -> bytes:
    document = Document()
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Aptos"
    normal_style.font.size = Pt(11)
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1)

    logo_paragraph = document.add_paragraph()
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if WORDMARK_PATH.exists():
        logo_paragraph.add_run().add_picture(str(WORDMARK_PATH), width=Inches(1.35))

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(22)
    title_run = title.add_run("Internship Completion Letter")
    set_run_font(title_run, bold=True, size=15, color="000000")

    date_paragraph = add_docx_paragraph(document, format_document_date(request.issued_date), after=18)
    date_paragraph.paragraph_format.space_after = Pt(18)

    company = document.add_paragraph()
    company.paragraph_format.space_after = Pt(22)
    company.paragraph_format.line_spacing = 1.25
    company_name = company.add_run("ReKnew Business Solutions Inc.")
    set_run_font(company_name, bold=True)
    for line in ("\n7800 N. Dallas Pkwy, Ste. 320", "\nPlano, TX 75024"):
        run = company.add_run(line)
        set_run_font(run)

    for paragraph_text in internship_completion_body_parts(request):
        add_docx_paragraph(document, paragraph_text, after=12)

    if SIGNATURE_OFFICIAL_PATH.exists():
        signature = document.add_paragraph()
        signature.paragraph_format.space_before = Pt(0)
        signature.paragraph_format.space_after = Pt(0)
        signature.add_run().add_picture(str(SIGNATURE_OFFICIAL_PATH), width=Inches(1.55))

    signer = document.add_paragraph()
    signer.paragraph_format.space_after = Pt(0)
    signer_run = signer.add_run("Murali Sajja | CEO")
    set_run_font(signer_run, bold=True, size=10, color="000000")

    company_signoff = document.add_paragraph()
    company_signoff.paragraph_format.space_after = Pt(0)
    signoff_run = company_signoff.add_run("ReKnew | reknew.ai")
    set_run_font(signoff_run, bold=True, size=10, color="000000")

    footer = section.footer
    footer.paragraphs[0].text = ""
    rule = footer.paragraphs[0]
    rule.paragraph_format.space_after = Pt(4)
    rule_run = rule.add_run("_" * 110)
    set_run_font(rule_run, size=6, color="000000")

    footer_line = footer.add_paragraph()
    footer_line.paragraph_format.space_before = Pt(0)
    footer_line.paragraph_format.space_after = Pt(0)
    left_run = footer_line.add_run("© ReKnew. All Rights Reserved.")
    set_run_font(left_run, size=7.5, color="8A8A8A")
    tab_run = footer_line.add_run("\t")
    set_run_font(tab_run, size=7.5, color="8A8A8A")
    right_run = footer_line.add_run("www.reknew.ai")
    set_run_font(right_run, size=7.5, color="8A8A8A")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
